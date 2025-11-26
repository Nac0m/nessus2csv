
#!/usr/bin/env python3
import argparse
import csv
import os
from collections import Counter

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---- Table headers (exact order) ----
HEADERS = [
    "Compliance CIS Benchmark",
    "Compliance Info",
    "Result Value",
    "Result",
    "CIS Recommendation",
    "Management Respond",
    "Revalidation Result",
]

# ---- Mapping from CSV columns to Word table columns ----
CSV_MAP = {
    "Compliance CIS Benchmark": "Compliance Check Name",   # will be combined with Benchmark Profile
    "Compliance Info": "Compliance Info",
    "Result Value": "Compliance Actual Value",
    "Result": "Compliance Result",
    "CIS Recommendation": "Compliance Solution",
    "Management Respond": None,      # leave blank
    "Revalidation Result": None,     # leave blank
}

# ---- Color palette ----
COLOR_TEXT_PASS   = RGBColor(0x00, 0x7A, 0x35)  # green
COLOR_TEXT_FAIL   = RGBColor(0xD9, 0x2F, 0x2F)  # red
COLOR_TEXT_OTHER  = RGBColor(0xB3, 0x6B, 0x00)  # amber

COLOR_BG_PASS     = "E7F4EA"  # light-green
COLOR_BG_FAIL     = "FDEAEA"  # light-red
COLOR_BG_OTHER    = "FFF4E5"  # light-amber

# ---- Utilities ----
def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)

def add_heading(doc: Document, text: str, size: int = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def set_header_style(cell):
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    run.font.size = Pt(11)

def set_cell_background(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), hex_color)

def set_result_cell_style(cell, value: str):
    normalized = (value or "").strip().upper()
    if normalized == "PASSED":
        text_color = COLOR_TEXT_PASS
        bg_color = COLOR_BG_PASS
    elif normalized == "FAILED":
        text_color = COLOR_TEXT_FAIL
        bg_color = COLOR_BG_FAIL
    else:
        text_color = COLOR_TEXT_OTHER
        bg_color = COLOR_BG_OTHER

    cell.text = ""  # clear
    p = cell.paragraphs[0]
    run = p.add_run(value or "")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = text_color
    set_cell_background(cell, bg_color)

def read_csv_rows(csv_path: str):
    with open(csv_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        return reader.fieldnames, rows

def tally_results(rows):
    counts = Counter()
    for r in rows:
        val = (r.get("Compliance Result") or "").strip().upper()
        counts[val] += 1
    total = sum(counts.values())
    return counts, total

def add_summary_section(doc: Document, counts: Counter, total: int):
    add_heading(doc, "Summary", size=13)

    passed = counts.get("PASSED", 0)
    failed = counts.get("FAILED", 0)
    pass_rate = (passed / total * 100.0) if total else 0.0
    fail_rate = (failed / total * 100.0) if total else 0.0

    p = doc.add_paragraph()
    p.add_run(f"Total checks: {total}\n").bold = True
    p.add_run(f"Passed: {passed} ({pass_rate:.1f}%)\n")
    p.add_run(f"Failed: {failed} ({fail_rate:.1f}%)\n")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Result"
    hdr[1].text = "Count"
    hdr[2].text = "Percentage"
    for c in hdr:
        set_header_style(c)

    def sort_key(k):
        priority = {"PASSED": 0, "FAILED": 1}
        return (priority.get(k, 2), k)

    for result_value in sorted(counts.keys(), key=sort_key):
        count = counts[result_value]
        pct = (count / total * 100.0) if total else 0.0
        row = table.add_row().cells
        row[0].text = result_value if result_value else "(blank)"
        row[1].text = str(count)
        row[2].text = f"{pct:.1f}%"

    doc.add_paragraph()  # spacer

def autofit_table_to_page(table, doc):
    # Compute usable page width (Length object)
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    # Lock widths so Word doesn't auto-resize columns
    table.allow_autofit = False

    # Proportional widths for the 7 columns (sum ~ 1.0)
    proportions = [0.18, 0.22, 0.15, 0.12, 0.22, 0.055, 0.055]

    # Apply widths to all rows (header + data)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = int(usable_width * proportions[i])  # EMUs

def create_compliance_table(doc: Document, rows: list):
    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"

    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(HEADERS):
        hdr_cells[i].text = header
        set_header_style(hdr_cells[i])

    # Populate rows
    for r in rows:
        tr = table.add_row().cells
        for col_idx, header in enumerate(HEADERS):
            is_result_col = (header == "Result")
            csv_col = CSV_MAP.get(header)
            value = ""
            if csv_col:
                value = (r.get(csv_col, "") or "").strip()

            # Combine Check Name with (Benchmark Profile) for Benchmark column
            if header == "Compliance CIS Benchmark":
                check_name = (r.get("Compliance Check Name") or "").strip()
                profile = (r.get("Compliance Benchmark Profile") or "").strip()
                value = f"{check_name} ({profile})" if profile else check_name

            if is_result_col:
                set_result_cell_style(tr[col_idx], value)
            else:
                tr[col_idx].text = value

    # Fit table to page width
    autofit_table_to_page(table, doc)
    return table

def set_landscape(doc: Document):
    # Apply landscape to all sections and swap page size
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

def main():
    parser = argparse.ArgumentParser(
        description="Create a new Word (.docx) with a color-coded compliance table and a PASSED/FAILED summary (landscape)."
    )
    parser.add_argument("-i", "--input", required=True, help="Path to input CSV file.")
    parser.add_argument("-o", "--output", required=True, help="Path to output Word (.docx) file.")
    parser.add_argument("--title", default="Compliance Results", help="Optional document title.")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        raise FileNotFoundError(f"CSV not found: {args.input}")

    # Read CSV
    fieldnames, rows = read_csv_rows(args.input)

    # Validate required CSV columns
    required = [src for src in CSV_MAP.values() if src is not None] + ["Compliance Benchmark Profile"]
    missing = [col for col in required if col not in fieldnames]
    # 'Compliance Benchmark Profile' is optional for combination; if missing, we won't fail
    if "Compliance Benchmark Profile" in missing:
        missing.remove("Compliance Benchmark Profile")
    if missing:
        raise ValueError("CSV is missing required columns: " + ", ".join(missing))

    # Tally
    counts, total = tally_results(rows)

    # Create new document
    doc = Document()
    set_landscape(doc)

    add_title(doc, args.title)
    doc.add_paragraph()  # spacer

    # Summary
    add_summary_section(doc, counts, total)

    # Compliance table
    create_compliance_table(doc, rows)

    # Save
    doc.save(args.output)
    print(f"[+] Created '{args.output}' from '{args.input}'.")
    print(f"    Summary: PASSED={counts.get('PASSED', 0)}, FAILED={counts.get('FAILED', 0)}, TOTAL={total}")

if __name__ == "__main__":
    main()