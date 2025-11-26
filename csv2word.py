
#!/usr/bin/env python3
import argparse
import csv
import os
from collections import Counter

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Mm


# ===== Table headers (exact order) =====
HEADERS = [
    "Compliance CIS Benchmark",
    "Compliance Info",
    "Result Value",
    "Result",
    "CIS Recommendation",
    "Management Respond",
    "Revalidation Result",
]

# ===== Mapping from CSV columns to Word table columns =====
CSV_MAP = {
    "Compliance CIS Benchmark": "Compliance Check Name",
    "Compliance Info": "Compliance Info",
    "Result Value": "Compliance Actual Value",
    "Result": "Compliance Result",
    "CIS Recommendation": "Compliance Solution",
    "Management Respond": None,      # leave blank
    "Revalidation Result": None,     # leave blank
}

# ===== Color palette (editable) =====
COLOR_TEXT_PASS   = RGBColor(0x00, 0x00, 0x00)  # green
COLOR_TEXT_FAIL   = RGBColor(0x00, 0x00, 0x00)  # red
COLOR_TEXT_OTHER  = RGBColor(0x00, 0x00, 0x00)  # amber

COLOR_BG_PASS     = "008000"  # green
COLOR_BG_FAIL     = "FF0000"  # red
COLOR_BG_OTHER    = "FFF4E5"  # light-amber

# ===== Utilities =====
def add_title(doc: Document, text: str):
    """Add a centered, bold title to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)

def add_heading(doc: Document, text: str, size: int = 12):
    """Add a left-aligned section heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def set_header_style(cell):
    """Style header cells: bold, slightly larger font."""
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = True
    run.font.size = Pt(11)

def autofit_table(table):
    """Basic alignment and approximate widths (optional)."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [
        1.8,  # Compliance CIS Benchmark
        2.2,  # Compliance Info
        1.5,  # Result Value
        1.2,  # Result
        2.2,  # CIS Recommendation
        1.6,  # Management Respond
        1.6,  # Revalidation Result
    ]
    for col_idx, width in enumerate(widths):
        for row in table.rows:
            row.cells[col_idx].width = Inches(width)

def set_cell_background(cell, hex_color: str):
    """Set table cell background shading using Word's w:shd element."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), hex_color)

def set_result_cell_style(cell, value: str):
    """
    Apply color coding based on Result value.
    - PASSED: green text + green background
    - FAILED: red text + red background
    - otherwise: amber text + amber background
    """
    normalized = (value or "").strip().upper()
    if normalized == "PASSED":
        text_color = COLOR_TEXT_PASS
        bg_color = COLOR_BG_PASS
    elif normalized == "FAILED":
        text_color = COLOR_TEXT_FAIL
        bg_color = COLOR_BG_FAIL
    else:
        text_color = COLOR_TEXT_OTHER
        bg_color = COLOR_BG_OTHER

    cell.text = ""  # clear any existing plain text
    p = cell.paragraphs[0]
    run = p.add_run(value or "")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = text_color
    set_cell_background(cell, bg_color)

# ===== CSV parsing helpers =====
def read_csv_rows(csv_path: str):
    """Return DictReader and a cached list of rows."""
    with open(csv_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        return reader.fieldnames, rows

def tally_results(rows):
    """
    Count occurrences in 'Compliance Result' column.
    Returns Counter and total count.
    """
    counts = Counter()
    for r in rows:
        val = (r.get("Compliance Result") or "").strip().upper()
        counts[val] += 1
    total = sum(counts.values())
    return counts, total

# ===== Summary section =====
def add_summary_section(doc: Document, counts: Counter, total: int):
    """Add a textual summary + a small table showing counts & percentages per result."""
    add_heading(doc, "Summary", size=13)

    passed = counts.get("PASSED", 0)
    failed = counts.get("FAILED", 0)
    # Avoid ZeroDivisionError
    pass_rate = (passed / total * 100.0) if total else 0.0
    fail_rate = (failed / total * 100.0) if total else 0.0

    # Summary paragraph
    p = doc.add_paragraph()
    p.add_run(f"Total checks: {total}\n").bold = True
    p.add_run(f"Passed: {passed} ({pass_rate:.1f}%)\n")
    p.add_run(f"Failed: {failed} ({fail_rate:.1f}%)\n")

    # Summary table (Result | Count | Percentage)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Result"
    hdr[1].text = "Count"
    hdr[2].text = "Percentage"
    for c in hdr:
        set_header_style(c)

    # For stable display: sort keys, but ensure PASSED/FAILED first
    def sort_key(k):
        priority = {"PASSED": 0, "FAILED": 1}
        return (priority.get(k, 2), k)

    for result_value in sorted(counts.keys(), key=sort_key):
        count = counts[result_value]
        pct = (count / total * 100.0) if total else 0.0
        row = table.add_row().cells
        row[0].text = result_value if result_value else "(blank)"
        row[1].text = str(count)
        row[2].text = f"{pct:.1f}%"

    doc.add_paragraph()  # spacer after summary

# ===== Main compliance table =====
def create_table_from_rows(doc: Document, rows: list):
    """Create the compliance table and populate it from already-read rows."""
    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"

    # Header cells
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(HEADERS):
        hdr_cells[i].text = header
        set_header_style(hdr_cells[i])

    # Populate rows
    for r in rows:
        tr = table.add_row().cells
        for col_idx, header in enumerate(HEADERS):
            csv_col = CSV_MAP[header]
            is_result_col = (header == "Result")
            value = ""
            if csv_col:
                value = (r.get(csv_col, "") or "").strip()


# Combine Check Name with (Benchmark Profile) for Benchmark column
            if header == "Compliance CIS Benchmark":
                check_name = (r.get("Compliance Check Name") or "").strip()
                profile = (r.get("Compliance Benchmark Profile") or "").strip()
                value = f"{check_name} ({profile})" if profile else check_name


            if is_result_col:
                set_result_cell_style(tr[col_idx], value)
            else:
                tr[col_idx].text = value

    autofit_table(table)
    return table

# ===== Entrypoint =====
def main():
    parser = argparse.ArgumentParser(
        description="Create a new Word (.docx) with a color-coded compliance table and a PASSED/FAILED summary."
    )
    parser.add_argument(
        "-i", "--input", required=True,
        help="Path to input CSV file (e.g., nessus_results.csv)."
    )
    parser.add_argument(
        "-o", "--output", required=True,
        help="Path to output Word file to create (e.g., compliance_report.docx)."
    )
    parser.add_argument(
        "--title", default="Compliance Results",
        help="Optional document title (default: 'Compliance Results')."
    )

    args = parser.parse_args()

    if not os.path.exists(args.input):
        raise FileNotFoundError(f"CSV not found: {args.input}")

    # Read CSV once
    fieldnames, rows = read_csv_rows(args.input)

    # Validate required CSV columns
    required = [src for src in CSV_MAP.values() if src is not None]
    missing = [col for col in required if col not in fieldnames]
    if missing:
        raise ValueError("CSV is missing required columns: " + ", ".join(missing))

    # Build summary
    counts, total = tally_results(rows)

    # Create a brand-new document
    doc = Document()
    
 #start a new section on a fresh page and force landscape + explicit size
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

# optional: set margins explicitly to avoid layout conflicts
    section.left_margin   = Mm(15)
    section.right_margin  = Mm(15)
    section.top_margin    = Mm(15)

    add_title(doc, args.title)
    doc.add_paragraph()  # spacer

    # Add Summary section
    add_summary_section(doc, counts, total)

    # Add Compliance table
    create_table_from_rows(doc, rows)

    # Save
    doc.save(args.output)
    print(f"[+] Created new document '{args.output}' from CSV '{args.input}'.")
    print(f"    Summary: PASSED={counts.get('PASSED', 0)}, FAILED={counts.get('FAILED', 0)}, TOTAL={total}")

if __name__ == "__main__":
    main()
